from win32com import client as wc
from docx import Document
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage, PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams
from io import StringIO
import pdfplumber
import os
import cchardet
from bs4 import BeautifulSoup
import re
import math

suffix_check = ['.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.pdf', '.html', '.htm', '.shtml', '.shtm', '.zip', '.rar', '.tar', '.bz2', '.7z', '.gz']
suffix_file = ['.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.pdf', '.zip', '.rar', '.tar', '.bz2', '.7z', '.gz']


# 下面是office 2007支持的全部文件格式对应表：
# wdFormatDocument = 0
# wdFormatDocument97 = 0
# wdFormatDocumentDefault = 16
# wdFormatDOSText = 4
# wdFormatDOSTextLineBreaks = 5
# wdFormatEncodedText = 7
# wdFormatFilteredHTML = 10
# wdFormatFlatXML = 19
# wdFormatFlatXMLMacroEnabled = 20
# wdFormatFlatXMLTemplate = 21
# wdFormatFlatXMLTemplateMacroEnabled = 22
# wdFormatHTML = 8
# wdFormatPDF = 17
# wdFormatRTF = 6
# wdFormatTemplate = 1
# wdFormatTemplate97 = 1
# wdFormatText = 2
# wdFormatTextLineBreaks = 3
# wdFormatUnicodeText = 7
# wdFormatWebArchive = 9
# wdFormatXML = 11
# wdFormatXMLDocument = 12
# wdFormatXMLDocumentMacroEnabled = 13
# wdFormatXMLTemplate = 14
# wdFormatXMLTemplateMacroEnabled = 15
# wdFormatXPS = 18


def doc_to_docx(filepath):
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(filepath)
    docxname = filepath + 'x'
    doc.SaveAs(docxname, 12)  # 使用参数12，表示将doc转成docx
    doc.Close()
    word.Quit()


def pdf_to_string(path):
    output = StringIO()
    with open(path, 'rb') as f:
        parser = PDFParser(f)
        doc = PDFDocument(parser)
        if not doc.is_extractable:
            raise PDFTextExtractionNotAllowed
        pdfrm = PDFResourceManager()
        laparams = LAParams()
        device = PDFPageAggregator(pdfrm, laparams=laparams)
        interpreter = PDFPageInterpreter(pdfrm, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)
            layout = device.get_result()
            for x in layout:
                if hasattr(x, "get_text"):
                    content = x.get_text()
                    output.write(content)
    content = output.getvalue()
    output.close()
    return content


def extract_pdf_tables(path) -> list:
    with pdfplumber.open(path) as pdf:
        list_tables = []
        for page in pdf.pages:
            for table in page.extract_tables():
                list_rows = []
                for row in table:
                    row_content = []
                    for one in row:
                        if one:
                            one = one.replace('\n', '')
                        row_content.append(one)
                    # print(row_content)
                    list_rows.append(row_content)
                list_tables.append(list_rows)
                # print('---------- 分割线 ----------')
        return list_tables


def divide_pdf_tables_text(path):
    with pdfplumber.open(path) as pdf:
        list_tables = []
        text = ''
        for page in pdf.pages:
            tables = page.find_tables()
            words = page.extract_words()
            words_left = ''
            if tables:
                words_left_page = ''
                for table in tables:
                    table_content = []
                    rows = table.extract()  # 解析出一张表格中的所有内容，每行内容为一个集合
                    for row in rows:
                        row_content = []
                        # 如果每行第一个ceil是空值，则认为该行不是单独的行，需要并入上一行内容，进行处理；
                        # 如果该行为该页表格的第一行，又是空值，则需要（1）将该表格并入上一页的表格，（2）将该行并入上一页表格中的最后一行进行处理；

                        for ceil in row:
                            # 遍历每一行中的每一个ceil, 如果不为空，则删除其中空格后，添加进row_content集合
                            # 如果ceil为空，None，或者‘’，则保留，并放入row_content中，
                            # 为的是后续执行check_table_rows()方法时，是通过每行第一个单元格是否是None，来判断是否需要并入上一行的
                            if ceil:  # 需要在这里加个判断，以防止正则匹配NoneType的error出现
                                ceil = re.sub(r'\s+', '', ceil)
                            row_content.append(ceil)  # 允许每一行第一个ceil为空，添加入row_content集合，这样解析出来的表格的列数不会发生变化
                        table_content.append(row_content)
                    list_tables.append(table_content)

                    x0, top, x1, bottom = table.bbox
                    # 表格top数值应大于表格前面文字的最后一个字的bottom
                    # 表格bottom数值应小于表格后面文字的第一个字的top
                    for word in words:
                        if word:
                            if word['bottom'] < top:
                                words_left_page += word['text']
                            elif word['top'] > bottom:
                                words_left_page += word['text']
                words_left += words_left_page
            else:
                for word in words:
                    if word:
                        words_left += word['text']
            text += words_left

        list_tables = table_to_two_columns(list_tables)
        list_tables_checked = check_list_table(list_tables)
        list_tables_rows_checked, table_text_left = check_table_rows(list_tables_checked)

        text += table_text_left
        return list_tables_rows_checked, text


# 将PDF文件表格中的内容也转换为text
def divide_pdf_full_text(path):
    with pdfplumber.open(path) as pdf:
        tables_text = ''
        text = ''
        for page in pdf.pages:
            tables = page.find_tables()
            words = page.extract_words()
            if tables:
                for table in tables:
                    rows = table.extract()
                    for row in rows:
                        if row:
                            for ceil in row:
                                if ceil:
                                    ceil = re.sub(r'\s+', '', ceil)
                                    tables_text += ceil

                    x0, top, x1, bottom = table.bbox
                    for word in words:
                        if word:
                            if word['bottom'] < top:                # 表格前面文字的最后一个字的bottom数值 小于 表格的top的数值
                                text += word['text']
                            elif word['top'] > bottom:              # 表格后面文字的第一个字的top 大于 表格bottom数值
                                text += word['text']
            else:
                for word in words:
                    if word:
                        text += word['text']

        return [tables_text], text


def divide_word_tables_text(path):
    if path.endswith('.doc'):
        doc_to_docx(path)
        os.remove(path)
        path = path + 'x'

    doc = Document(path)
    list_tables = []
    for table in doc.tables:
        table_content = []
        for i, row in enumerate(table.rows):
            row_content = []
            cell_set = set()
            for cell in row.cells:
                if cell not in cell_set:  # 合并的单元格虽然重复但公用内存地址是相同的
                    cell_set.add(cell)
                    row_content.append(cell.text)
            table_content.append(row_content)
        list_tables.append(table_content)

    list_tables = table_to_two_columns(list_tables)
    list_tables_checked = check_list_table(list_tables)
    list_tables_rows_checked, table_text_left = check_table_rows(list_tables_checked)

    text = ''
    for p in doc.paragraphs:
        text += p.text.strip()

    text += table_text_left
    return list_tables_rows_checked, text


# 将多列表格转换为两列表格
def divide_html_tables_text(path):
    with open(path, 'rb') as f:
        content = f.read()
        encoding = cchardet.detect(content)['encoding']
        content = content.decode(encoding, errors='ignore')
        soup = BeautifulSoup(content, 'lxml')
        list_tables = []
        tables = soup.find_all(name='table')
        for table in tables:
            list_rows = []
            rows = table.find_all(name='tr')
            for row in rows:
                list_tds = []
                tds = row.find_all(name='td')
                for td in tds:
                    td_text = re.sub(r'\s+', '', td.text)
                    list_tds.append(td_text)
                list_rows.append(list_tds)
            list_tables.append(list_rows)

        list_tables = table_to_two_columns(list_tables)
        list_tables_checked = check_list_table(list_tables)
        list_tables_rows_checked, table_text_left = check_table_rows(list_tables_checked)

        re_table = re.compile('<\s*table[^>]*>.*?<\s*/\s*table\s*>', re.DOTALL)
        re_comment = re.compile('<!--[^>]*-->')
        content_left = re_table.sub('', content)
        content_left = re_comment.sub('', content_left)
        soup_left = BeautifulSoup(content_left, 'lxml')
        text = soup_left.text

        text += table_text_left
        return list_tables_rows_checked, text


# 将6列以上的表格行转换为两列表格
def table_to_two_columns(list_tables: list):
    list_tables_new = []
    for table in list_tables:
        table_content = []
        row_header = []
        for row in table:
            if len(row) > 6:
                # 如果该行中的列数大于6，并且row_header为空集合，则将该行设置为row_header
                if len(row_header) == 0:
                    row_header = row
                else:
                    # 如果该行中的列数大于6，并且row_header已经有内容了，则该行为row_body
                    row_body = row
                    # 只有在row_header和row_body的列数相等的情况下，才做行列转换，
                    # 并且转换后，作为一张新的表格内容table_content_new放入list_tables_new集合中
                    if len(row_header) == len(row_body):
                        table_content_new = []
                        for i in range(len(row_header)):
                            row_new = [row_header[i], row_body[i]]
                            table_content_new.append(row_new)
                        list_tables_new.append(table_content_new)
                    else:
                        # 否则，将row_header和row_body都加入table_content集合中后，将row_header恢复为空集合
                        table_content.append(row_header)
                        table_content.append(row_body)
                        row_header = []
            else:
                # 如果每一行中的列数小于等于6，则该行直接加入table_content集合中
                # 并且将row_header回复为空集合
                table_content.append(row)
                row_header = []
        list_tables_new.append(table_content)
    return list_tables_new


# 转换具有合并行，和 合并列的表格，到一个list_cells集合中index=[行号数字，列号数字]
def table_to_cells(table):
    list_cells = []
    row_cells, column_cells = [], []
    index = []
    width, length = len(table.columns), len(table.rows)
    k = 0
    for row in table.rows:
        for cell in row.cells:
            if cell not in row_cells:
                index.append([k // width, k % width])
                row_cells.append(cell)
            k += 1
    k = 0
    for column in table.columns:
        for cell in column.cells:
            if cell not in column_cells:
                column_cells.append(cell)
            elif [k % length, k // length] in index:
                index.remove([k % length, k // length])
            k += 1
    for i in index:
        text = table.rows[i[0]].cells[i[1]].text
        list_cells.append(text)
    return list_cells


# 此处开始检查每一页的表格是否是完整的独立表格，或者是连接上下页的表格
# 通过关键字检查，来区分是否为一张独立的表格，关键字为“产品名称”，“产品代码”或“产品编码”
# 将表格每行，转换为2个单元格的结构，即row[0]为label, row[1]为value
def check_list_table(list_tables):
    list_tables_checked = []
    if len(list_tables) == 0:
        return list_tables_checked

    for index, table in enumerate(list_tables):
        # 兴业银行：期次款数；江苏银行：子份额
        list_product_name = ['名称', '子份额', '期次款数']
        list_product_code = ['代码', '编码', '编号']
        name_in = False
        code_in = False
        table_content = []
        for row in table:
            if row:  # 此处过滤掉空集合的row
                label = row[0]  # 此处没有检查label是否为空值
                value = row[1] if len(row) > 1 else ''
                value = value or ''  # 去除value等于None的结果，将None替换为''值
                # !!!! 如果表格列数超过2列，则把后面列的单元格里的内容，加在第二个单元格内容中，以空格区分
                if len(row) > 2:
                    for i in range(2, len(row)):
                        if row[i]:
                            value += ('\t' + row[i])

                row_content = [label, value]
                table_content.append(row_content)

                if label:
                    for name in list_product_name:
                        if name in label:
                            name_in = True
                            break
                    for code in list_product_code:
                        if code in label:
                            code_in = True
                            break
                if value:
                    for code in list_product_code:
                        if code in value:
                            code_in = True
                            break

        # 只有同时存在产品名称和产品编号的table，可以直接添加进list_tables_checked集合
        if name_in and code_in:
            list_tables_checked.append(table_content)
        else:
            # 否则，从list_tables_checked集合中取出最后一个合格的table，加在后面
            if len(list_tables_checked) > 0:
                last_table = list_tables_checked[-1]
                list_tables_checked = list_tables_checked[:-1]

                last_table += table_content  # last_table合并当前的table集合
                list_tables_checked.append(last_table)  # 再重新添加进list_tables_checked集合
            else:
                # ！！！如果list_tables_checked还是空集合
                # if name_in or code_in:                        # 是否要设置条件：允许放入一张只含有产品名称或只含有产品编号的表格作为第一张表格 ？？？
                list_tables_checked.append(table_content)  # 如果list_tables_checked还是空集合，则直接添加表格进去
    return list_tables_checked


def check_table_rows(list_tables):
    list_table_checked = []
    table_text_left = ''
    if len(list_tables) == 0:
        return list_table_checked

    for table in list_tables:
        table_empty_label_checked = []
        for row in table:
            if row:
                label = row[0]
                value = row[1]
                if label:
                    if len(label) < 21:  # 如果label为20个字以内的，则正常添加进row
                        table_empty_label_checked.append(row)
                    else:
                        if len(table_empty_label_checked) > 0:
                            last_row = table_empty_label_checked[-1]  # 否则取出已经检查的table中的，最后1行内容
                            table_empty_label_checked = table_empty_label_checked[:-1]  # 并从已经检查的table中，去除这最后1行

                            last_row_label = last_row[0]
                            last_row_value = last_row[1]
                            if last_row_label and not last_row_value:  # last_row的label有值，而value为空值
                                last_row_value = label + '\t' + value  # 将label和value的值合并后，设置为last_row_value的值
                                row_new = [last_row_label, last_row_value]  # 重新创建一个新的label和value的list的row
                                table_empty_label_checked.append(row_new)  # 将这个新的row加回到已经检查的table中去
                            else:
                                table_text_left += (label + '\t' + value)  # 如果last_row的label有值，而value也有值
                        else:
                            table_text_left += (label + '\t' + value)  # 不管value有没有值，跟label合并后，添加到table_text_left当中去
                else:
                    if len(table_empty_label_checked) > 0:  # 如果已经检查的table中，至少存在1行
                        if value:  # 如果value值 不为空
                            last_row = table_empty_label_checked[-1]  # 则取出已经检查的table中的，最后1行内容
                            table_empty_label_checked = table_empty_label_checked[:-1]  # 并从已经检查的table中，去除这最后1行

                            last_row_label = last_row[0]
                            last_row_value = last_row[
                                                 1] or ''  # 避免出现None值，下面+=运算出现， unsupported operand type(s) for +=: 'NoneType' and 'str'的错误

                            last_row_value += value  # 把取出的这最后一行的value，加上待添加的value
                            row_new = [last_row_label, last_row_value]  # 后重新创建一个新的label和value的list的row
                            table_empty_label_checked.append(row_new)  # 将这个新的row加回到已经检查的table中去
                    else:  # 如果label是空值，但是value不是空值，而且table_empty_label_checked集合中还没有添加任何row进去
                        if value:  # 因为后面解析表格，是按照label来匹配的，如果没有label，则value也会被忽略掉
                            table_text_left += value  # 将这部分多出来的value, 以table_text_left的形式汇总后，返回出去
        list_table_checked.append(table_empty_label_checked)
    return list_table_checked, table_text_left
